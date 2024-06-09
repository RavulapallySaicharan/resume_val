import PyPDF2
import docx
import json
from openai import OpenAI
import psycopg2
from sqlalchemy import create_engine

CANDIDATE_TABLE = "candidate_table_test_entire_summary123"
VECTOR_TABLE = "vector_table_test_entire_summary123"
JD_TABLE = "jd_candidate_table_test_entire_summary123"


API_KEY = "sk-<>"
MODEL = "gpt-3.5-turbo"
MODEL_4o = "gpt-4o"
EMBEDDINGS_MODEL = "text-embedding-3-large"

HOST = 'database-1.cvo0sykm8au5.us-east-2.rds.amazonaws.com'
PORT = '5432'  # default PostgreSQL port is 5432
DATABASE = 'first_simple_db'
USER = 'postgres'
PASSWORD = 'CharanKranthi1!'



# def create_resume_extraction_prompt(text):
#     prompt_template = f"""

#     Below is a resume. You are a helpfull assistant to read the complete
#     resume and extract relevant infromation which can be used for understanding the candidate skills and other
#     details which are helpfull for recruiters
    
#     From the resume extract information for the following keywords ['name','location', 'email', 'mobile number',
#     'college', 'designation', 'total experience', 'list of companies worked', 'list of skillset',
#     'degree', 'list of domains candidate has worked on', 'summary']. Identify these details from the entire resume. Make sure to not miss any details.
#     name : Extract Name of the candidate from the entire resume.
#     location : Extract Location of the candidate from the entire resume.
#     email : Extract email of the candidate from the entire resume.
#     mobile number : Extract mobile number of the candidate from the entire resume.
#     Extract total experience in the format 5.1 years.
#     While extracting list of skillset include all the skills from the resume.
#     Write a detailed 500-word summary explaining the details.Make sure the summary should prioritize the following elements in order: Job Designation, Experience in '5.1 years' format, 
#     Skills needed, brief description of the candidates responsibilites in projects, Highest degree, Location, Hourly rate, notice period, List of certifications done by the candidate.
#     Identify these details from the entire resume. Make sure to not miss any details.
#     Use the below summary format to generate.
#     Job Designation: . Experience: 5.1 years. Skills: .Responsibilities: . Highest degree: . Location: .
#     Hourly rate: . Notice period: . List of certifications: .
    
#     Your output should be in a json format as described below, including the leading and trailing "json" and "":
#     The output json should not include any characters which are not supported by json. 
#     Return the output in correct json format. Do not include comma after summary. Do not include any new key into the json. Use only below given keys.

#         'name': string,
#         'location': string,
#         'email': string,
#         'mobile number': string,
#         'college': string,
#         'designation': string,
#         'total experience': string,
#         'list of companies worked': string,
#         'list of skillset': string,
#         'degree': string,
#         'list of domains candidate has worked on': string,
#         'summary': string

        
#     Here is the resume
#     {text}

#     If any details are not captured in the json, updated them as 'Not Mentioned'
#     """
#     return prompt_template

def create_resume_extraction_prompt(text):
    prompt_template = f"""

    Below is a resume. You are a helpfull assistant to read the complete
    resume and extract relevant infromation which can be used for understanding the candidate skills and other
    details which are helpfull for recruiters
    
    From the resume extract information for the following keywords ['name','location', 'email', 'mobile number',
    'college', 'designation', 'total experience', 'list of companies worked', 'list of skillset',
    'degree', 'list of domains candidate has worked on',hourlyrate, noticeperiod, 'summary']. Identify these details from the entire resume. Make sure to not miss any details.
    
    
    Your output should be in a json format as described below, including the leading and trailing "```json" and "```":
    The output json should not include any characters which are not supported by json. 
    Return the output in correct json format. Do not include comma after summary. Do not include any new key into the json. Use only below given keys.

        'name': Extract Name of the candidate from the entire resume,
        'location' : Extract the location of the candidate from the entire resume,
        'email': Extract email of the candidate from the entire resume,
        'mobileNumber': Extract mobile number of the candidate from the entire resume,
        'college': Extract college name of the candidate from the entire resume,
        'designation': Extract designation of the candidate from the entire resume,
        'totalExperience': Extract total experience in the format 5.1 years,
        'companiesWorked': Extract all the companies that the candidate has worked and create a list,
        'skillset': Extract all the technical skills from the candidate resume,
        'degree': Extract the highest degree of the candidate from the entire resume,
        'domians': Extract all the domains the cadidate has worked on  and create a list,
        'hourlyrate' : Extract the hourly rate of the candicate if mentioned in the  resume,
        'noticeperiod' : Extract the notice period of the candidate if mentioned in the resume,
        'summary': Analyze the entire resume and write a complete summary. Do not miss any details.
         
        
    Here is the resume
    {text}

    If any details are not captured in the json, updated them as 'Not Mentioned'
    """
    return prompt_template

def create_jd_summary_prompt(text):
    prompt_template = f"""

    Below is a Job description. You are a helpfull assistant to read the complete Job description and summarize infromation which can be used for understanding.
    
    Also write a detailed 500-word summary explaining the candidate's profile with his details. Make sure not to include the candidate's name in the summary. The summary should prioritize the following elements in order. Use this format to generate summary. Job Designation : . Relevant Experience : . Skills : . Roles and Responsibilities : Extract all the roles and responsiblities exhibited by the candiade from all his projects. Education : . Certifications: . If any details are not found use Not mentioned.

    Your output should be in a json format as described below, including the leading and trailing "json" and "":

        'summary': string

    {text}
    """

    return prompt_template


def get_ratings_prompt(resume, summary):
    prompt_template = f"""

    To assess if the below uploaded resume of the candidate matches the provided job description (JD), we'll evaluate it based on three main criteria: Skills, Experience, and Location. Read the entire resume to get the location. If any location is not mentioned then use 0.
    
    Below is the Job Description : 
    {summary}
    
    Below is the Resume:
    {resume}
    Compare the Job description against the resume and rate everything out of 10.


    Your output should be in a markdown format and strictly use below template.

    # Resume Assessment

    ## Job Description:
    - **Job Designation**: 
    - **Relevant Experience**: Compare the cadidate experience with job descrition experience (Rating).
    - **Skills**: Compare the candidate skills and job description skills. Include all the skills from Job description and resume (Rating).
    - **Location**: Compare the location of the cadidate with the location of the job description (Rating).
    ## Give a final ating out of 10.

    ## Summary
    - **Summary of observations.***
    
    """
    return prompt_template
    
def create_jd_creation_prompt(text):
    prompt_template = f"""
    I want you to act as a Job Description Generator. Your task is to generate a detailed job description based on the format provided below and the information given by the user. If any required information is missing or unclear from the user's input, please indicate it as 'ANY' instead of making assumptions.

    Your output should be in a markdown format.

    ---

    Job Title: [Enter Job Title Here]

    Job Location: [Enter Job Location Here]

    Job Type: [Full-time/Part-time/Contract]

    Job Summary:
    [Provide a brief overview of the role and its purpose within the organization.]

    Key Responsibilities:
        [List specific duties and tasks the candidate will be responsible for]
        [List additional responsibilities as necessary]

    Requirements:
        [Education level required]
        [Years of experience]
        [Specific skills or certifications required]
        [Soft skills such as communication, teamwork, etc.]

    Preferred Qualifications:
        [Additional qualifications or experiences that would be beneficial but not mandatory]

    ---

    Here is the user input:
    {text}
    """
    return prompt_template


def generate_response(prompt_template):

    client = OpenAI(api_key=API_KEY)

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "user",
                "content": prompt_template,
            }
            ]

            )
    return response

def generate_response_4o(prompt_template):

    client = OpenAI(api_key=API_KEY)

    response = client.chat.completions.create(
        model=MODEL_4o,
        messages=[
            {
                "role": "user",
                "content": prompt_template,
            }
            ]

            )
    return response

def convert_to_json(input_string):
    cleaned_string = input_string.replace('```', '').replace('json', '')
    # cleaned_string = cleaned_string.replace('\"','')
    data = json.loads(cleaned_string)

    return data

def get_embedding(text):
    # text = text.replace("\n", " ")
    client = OpenAI(api_key=API_KEY)
    return client.embeddings.create(input = [text], model=EMBEDDINGS_MODEL, dimensions = 1000).data[0].embedding


def read_documents(filename):
    if filename.endswith('.pdf'):
        with open(f'{filename}', 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            text = ""
            for i in range(num_pages):
                page = pdf_reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    non_empty_lines = [line for line in page_text.split('\n') if line.strip() != '']
                    text += '\n'.join(non_empty_lines) + '\n'

    if filename.endswith('.docx'):
        doc = docx.Document(f'{filename}')
        # doc = docx.Document(f'./ResumeData/{filename}')
        text = ""
        
        # Read headers and footers from all sections
        for section in doc.sections:
            # Read headers
            header = section.header
            for paragraph in header.paragraphs:
                text += paragraph.text + "\n"
                
            # Read footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                text += paragraph.text + "\n"
        
        # Read main content
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text += paragraph.text + "\n"
        
        
        # print(text)

    return text.replace('\0', '')

def get_connection():
    try:
        # Connect to the PostgreSQL database
        connection = psycopg2.connect(
            host=HOST,
            port=PORT,
            database=DATABASE,
            user=USER,
            password=PASSWORD
        )

        # Create a cursor to perform database operations
        cursor = connection.cursor()
        # record = cursor.fetchone()
        # print("You are connected to - ", record)
        engine = create_engine(f'postgresql://{USER}:{PASSWORD}@{HOST}:{PORT}/{DATABASE}')


    except (Exception, psycopg2.Error) as error:
        print("Error while connecting to PostgreSQL", error)

    return cursor, engine, connection

def exec_query(cursor, query, params=None):
    if params:
        cursor.execute(query, params)
    else:
        cursor.execute(query)

def get_records(cursor, table):
    select_query = f"SELECT * FROM {table};"
    cursor.execute(select_query)
    records = cursor.fetchall()

    # Print all records
    for record in records:
        print(record)

